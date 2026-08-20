"""Tests file extraction for every supported format and the /api/demo/upload endpoint
that uses it. CSV/Excel rows must become the REAL dataset (no synthesis at all);
docx/pdf/pptx text must flow through the same Requirement Agent as typed text."""
import io

import pytest
from fastapi.testclient import TestClient

from main import app
from services.file_import import (
    FileParseError,
    UnsupportedFileTypeError,
    extract_from_file,
)


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        yield c


@pytest.fixture(scope="module")
def client_headers(client):
    resp = client.post("/api/client/session", json={"email": "file-import-tests@example.com"})
    assert resp.status_code == 200
    return {"X-Client-Token": resp.json()["client_token"]}


def test_csv_extraction_is_tabular_with_real_rows():
    csv_bytes = b"customer_name,risk_score,status\nAcme Co,72,active\nBeta Inc,15,active\n"
    result = extract_from_file("customers.csv", csv_bytes)
    assert result.kind == "tabular"
    assert result.fields == ["customer_name", "risk_score", "status"]
    assert len(result.records) == 2
    assert result.records[0]["customer_name"] == "Acme Co"
    assert result.record_label == "customer"  # singularized from filename


def test_csv_numeric_columns_are_coerced_not_left_as_strings():
    """Regression: csv.DictReader returns every cell as a string. Without coercion,
    a numeric field like risk_score="72" silently fails every CHECK_CONDITION
    comparison (TypeError caught, treated as unmatched) - the bug this guards
    against was caught via a live browser test, not by static review."""
    csv_bytes = b"customer_name,risk_score,active\nAcme Co,72,true\nBeta Inc,15,false\n"
    result = extract_from_file("customers.csv", csv_bytes)
    assert result.records[0]["risk_score"] == 72
    assert isinstance(result.records[0]["risk_score"], int)
    assert result.records[0]["active"] is True
    assert result.records[1]["active"] is False


def test_xlsx_extraction_is_tabular_with_real_rows():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["product", "stock_level"])
    ws.append(["Widget", 12])
    ws.append(["Gadget", 88])
    buf = io.BytesIO()
    wb.save(buf)

    result = extract_from_file("inventory.xlsx", buf.getvalue())
    assert result.kind == "tabular"
    assert result.fields == ["product", "stock_level"]
    assert len(result.records) == 2
    assert result.records[0]["stock_level"] == 12


def test_docx_extraction_is_text():
    from docx import Document

    doc = Document()
    doc.add_paragraph("We want to track field sales visits and flag overdue ones.")
    buf = io.BytesIO()
    doc.save(buf)

    result = extract_from_file("brief.docx", buf.getvalue())
    assert result.kind == "text"
    assert "field sales visits" in result.text


def test_unsupported_extension_rejected():
    with pytest.raises(UnsupportedFileTypeError):
        extract_from_file("malware.exe", b"whatever")


def test_empty_csv_rejected():
    with pytest.raises(FileParseError):
        extract_from_file("empty.csv", b"")


def test_upload_csv_uses_real_rows_not_synthesis(client, client_headers):
    csv_bytes = b"customer_name,risk_score,status\nAcme Co,72,active\nBeta Inc,15,active\nGamma LLC,90,active\n"
    resp = client.post(
        "/api/client/demo/upload",
        files={"file": ("customers.csv", csv_bytes, "text/csv")},
        data={"instruction": "Flag customers with risk_score above 50 and alert the sales team."},
        headers=client_headers,
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["outcome"] == "executed"
    assert body["requirement"]["record_label"] == "customer"
    assert body["requirement"]["fields"] == ["customer_name", "risk_score", "status"]
    # the actual uploaded values must appear in the real per-step output
    assert "Acme Co" in str(body["execution"])
    # and the numeric threshold check must actually match (72 and 90 are > 50) -
    # this fails if risk_score wasn't coerced out of its raw CSV string form
    decision_step = next(s for s in body["execution"]["step_results"] if s["tool"] == "MAKE_DECISION")
    assert decision_step["output"]["decision"] == "flag_for_action"


def test_upload_docx_runs_through_same_pipeline_as_typed_text(client, client_headers):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Track field sales rep visits and flag reps overdue for a visit.")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        "/api/client/demo/upload",
        files={"file": ("brief.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=client_headers,
    )
    assert resp.status_code == 200
    assert resp.json()["outcome"] in ("executed", "blueprint")


def test_upload_rejects_unsupported_extension(client, client_headers):
    resp = client.post(
        "/api/client/demo/upload",
        files={"file": ("virus.exe", b"nope", "application/octet-stream")},
        headers=client_headers,
    )
    assert resp.status_code == 415


def test_upload_rejects_oversized_file(client, client_headers):
    big = b"a" * (5 * 1024 * 1024 + 1)
    resp = client.post(
        "/api/client/demo/upload", files={"file": ("big.csv", big, "text/csv")}, headers=client_headers
    )
    assert resp.status_code == 413
